import argparse
from PIL import Image
import pytesseract
from docx import Document

def extract_text_from_image(image_path):
    image = Image.open(image_path)
    text = pytesseract.image_to_string(image)

    return text

def create_word_document(text, output_path):
    document = Document()
    document.add_heading('Texto Extraído', level=1)

    # Limpiar el texto para que sea XML compatible
    clean_text = ''.join(char for char in text if 32 <= ord(char) <= 126 or ord(char) == 9 or ord(char) == 10 or ord(char) == 13)
    
    document.add_paragraph(clean_text)

    document.save(output_path)

def main():
    parser = argparse.ArgumentParser(description='Convertir imagen a texto y exportar a Word.')
    parser.add_argument('image_path', help='Ruta del archivo de imagen')
    args = parser.parse_args()

    # Extraer texto utilizando OCR
    image_text = extract_text_from_image(args.image_path)
    print("Texto extraído de la imagen:")

    # Obtener el nombre del archivo sin extensión
    file_name_without_extension = args.image_path.split('.')[0]

    # Exportar a un archivo Word con el mismo nombre del archivo de imagen
    output_word_file = f'{file_name_without_extension}.docx'
    create_word_document(image_text, output_path=output_word_file)
    print(f"\nTexto exportado a {output_word_file}")

if __name__ == "__main__":
    main()

