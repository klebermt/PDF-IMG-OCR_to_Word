import argparse
import fitz  # PyMuPDF
import pytesseract
from docx import Document

def convert_pdf_to_text(pdf_path):
    pdf_document = fitz.open(pdf_path)
    text = ""

    for page_num in range(pdf_document.page_count):
        page = pdf_document[page_num]
        text += page.get_text()

    pdf_document.close()

    return text

def extract_text_from_image(image_path):
    image = Image.open(image_path)
    text = pytesseract.image_to_string(image)

    return text

def create_word_document(text, output_path='output.docx'):
    document = Document()
    document.add_heading('Texto Extraído', level=1)
    document.add_paragraph(text)

    document.save(output_path)

def main():
    parser = argparse.ArgumentParser(description='Convertir PDF escaneado a texto y exportar a Word.')
    parser.add_argument('pdf_path', help='Ruta del archivo PDF escaneado')
    args = parser.parse_args()

    pdf_text = convert_pdf_to_text(args.pdf_path)
    print("Texto extraído del PDF escaneado:")

    # Obtener el nombre del archivo sin extensión
    file_name = args.pdf_path.split('.')[0]

    # Exportar a un archivo Word
    output_word = f'{file_name}.docx'
    create_word_document(pdf_text, output_path=output_word)
    print(f"\nTexto exportado a {output_word}")

if __name__ == "__main__":
    main()

