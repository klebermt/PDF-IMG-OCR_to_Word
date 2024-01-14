import argparse
import os
import fitz  # PyMuPDF
from PIL import Image
import pytesseract
from docx import Document
import string

def convert_pdf_to_images(pdf_path, output_folder='images'):
    pdf_document = fitz.open(pdf_path)
    image_paths = []

    # Crear la carpeta de salida si no existe
    os.makedirs(output_folder, exist_ok=True)

    for page_num in range(pdf_document.page_count):
        page = pdf_document[page_num]
        image_path = os.path.join(output_folder, f'page_{page_num + 1}.png')
        image_paths.append(image_path)

        # Obtener la imagen como un objeto Pixmap
        pixmap = page.get_pixmap()

        # Convertir el Pixmap a una imagen de Pillow
        image = Image.frombytes("RGB", [pixmap.width, pixmap.height], pixmap.samples)

        # Guardar la imagen en formato PNG
        image.save(image_path, "PNG")

    pdf_document.close()

    return image_paths

def extract_text_from_image(image_path):
    image = Image.open(image_path)
    text = pytesseract.image_to_string(image)

    return text

def clean_text(text):
    # Eliminar caracteres no imprimibles y de control
    cleaned_text = ''.join(char if char in string.printable else ' ' for char in text)
    return ' '.join(cleaned_text.split())

def create_word_document(text, output_path):
    document = Document()
    document.add_heading('Texto Extraído', level=1)

    # Limpiar el texto para que sea XML compatible
    clean_text_result = clean_text(text)
    document.add_paragraph(clean_text_result)

    document.save(output_path)

def main():
    parser = argparse.ArgumentParser(description='Convertir PDF a Word mediante OCR.')
    parser.add_argument('pdf_path', help='Ruta del archivo PDF')
    args = parser.parse_args()

    # Convertir PDF a imágenes
    image_paths = convert_pdf_to_images(args.pdf_path)
    print(f"Se han convertido {len(image_paths)} páginas del PDF a imágenes.")

    # Aplicar OCR y extraer texto de cada imagen
    text_per_page = []
    for image_path in image_paths:
        text = extract_text_from_image(image_path)
        text_per_page.append(text)

    # Concatenar el texto de todas las páginas
    all_text = '\n\n'.join(text_per_page)

    # Obtener el nombre del archivo PDF sin extensión
    file_name_without_extension = os.path.splitext(os.path.basename(args.pdf_path))[0]

    # Exportar a un archivo Word con el mismo nombre del archivo PDF
    output_word_file = f'{file_name_without_extension}_output.docx'
    create_word_document(all_text, output_path=output_word_file)
    print(f"\nTexto exportado a {output_word_file}")

if __name__ == "__main__":
    main()

